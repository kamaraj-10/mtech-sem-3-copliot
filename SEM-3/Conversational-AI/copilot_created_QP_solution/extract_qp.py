import os
from pathlib import Path
import PyPDF2
from docx import Document

# Get the directory containing this script
script_dir = Path(__file__).parent

# Extract from PDF (Sample Solution)
print("=" * 80)
print("EXTRACTING FROM PDF: ConvAI-Midsem_SampleQP-Solution.pdf")
print("=" * 80)

pdf_file = script_dir / "ConvAI-Midsem_SampleQP-Solution.pdf"
if pdf_file.exists():
    with open(pdf_file, 'rb') as file:
        pdf_reader = PyPDF2.PdfReader(file)
        print(f"\nTotal pages: {len(pdf_reader.pages)}\n")
        
        # Extract all text
        full_text = ""
        for i in range(len(pdf_reader.pages)):
            text = pdf_reader.pages[i].extract_text()
            full_text += f"\n--- PAGE {i+1} ---\n{text}"
        
        print(full_text)

# Extract from DOCX (Main QP)
print("\n\n" + "=" * 80)
print("EXTRACTING FROM DOCX: 2025-2026 QP")
print("=" * 80)

docx_file = script_dir / "2025-2026_SEM_1_AIML_AA05__AB05_1-2024_AIMLCZG521_CONVERSATIONAL_AI_EC2_REGULAR_21-12-2025_AN.docx"
if docx_file.exists():
    doc = Document(docx_file)
    print(f"\nTotal paragraphs: {len(doc.paragraphs)}\n")
    
    # Extract all text
    for p in doc.paragraphs:
        if p.text.strip():
            print(p.text)
